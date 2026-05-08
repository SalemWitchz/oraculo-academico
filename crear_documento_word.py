# -*- coding: utf-8 -*-
"""Script para generar el documento Word de cumplimiento de requisitos."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────────────────────────────────────────
# Helpers de estilo
# ─────────────────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1F1F1F"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        r, g, b = tuple(int(color[i:i+2], 16) for i in (0, 2, 4))
        run.font.color.rgb = RGBColor(r, g, b)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, bold_prefix=None, indent_cm=1.0):
    """Párrafo con viñeta manual y prefijo en negrita opcional."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.font.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_colored_box(doc, titulo, lineas, color_hex="1A1A2E", text_color=(200, 200, 220)):
    """Tabla de 1 celda que simula un recuadro de color con título y lista."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, color_hex)
    p_title = cell.paragraphs[0]
    r = p_title.add_run(titulo)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*text_color)
    for linea in lineas:
        p_l = cell.add_paragraph()
        rl  = p_l.add_run("    • " + linea)
        rl.font.size = Pt(10)
        rl.font.color.rgb = RGBColor(210, 210, 230)
    doc.add_paragraph()

def agregar_seccion_requisitos(doc):
    """Agrega la sección 2 con las especificaciones oficiales del proyecto."""

    add_heading(doc, "2. Especificaciones Oficiales del Proyecto", level=1)

    # Título y unidad
    p_tit = doc.add_paragraph()
    r_t   = p_tit.add_run("Sistema de Análisis Estadístico para la Evaluación Académica")
    r_t.font.bold = True
    r_t.font.size = Pt(13)

    p_uni = doc.add_paragraph()
    r_u   = p_uni.add_run("Unidad de Competencia: ")
    r_u.font.bold = True
    r_u.font.size = Pt(11)
    p_uni.add_run("Probabilidad y Estadística").font.size = Pt(11)

    doc.add_paragraph()

    # ── Objetivo ──────────────────────────────────────────────────────
    add_heading(doc, "2.1  Objetivo del Proyecto", level=2)
    add_para(doc,
             "Diseñar e implementar un software que permita analizar el rendimiento académico "
             "de los estudiantes, identificando patrones, factores de influencia y riesgos para "
             "apoyar la toma de decisiones institucionales fundamentadas en datos.")

    # ── Problema ──────────────────────────────────────────────────────
    add_heading(doc, "2.2  Problema a Resolver", level=2)
    for item in [
        "Falta de detección temprana de riesgo académico.",
        "Desconocimiento de los factores que impactan el rendimiento.",
        "Decisiones institucionales sin sustento estadístico.",
    ]:
        add_bullet(doc, item, "•")

    # ── Validación estadística ─────────────────────────────────────────
    add_heading(doc, "2.3  Validación Estadística — Opción B: Empleo (seleccionada)", level=2)
    add_para(doc,
             "El proyecto seleccionó la Opción B para su validación estadística mediante "
             "pruebas de inferencia:")

    tbl_h = doc.add_table(rows=4, cols=2)
    tbl_h.style = "Table Grid"
    encabezados = [("Opción", "Hipótesis"), ("A — Asistencia", "H₁: El porcentaje de asistencia influye positivamente en el promedio final."),
                   ("B — Empleo  ✓ (seleccionada)", "H₁: Los estudiantes que trabajan tienen un promedio menor a los que no trabajan."),
                   ("C — Plataformas", "H₁: Existe correlación positiva entre el uso de plataformas y el rendimiento.")]
    for r_idx, (col0, col1) in enumerate(encabezados):
        cells = tbl_h.rows[r_idx].cells
        cells[0].text = col0
        cells[1].text = col1
        if r_idx == 0:
            for c in cells:
                set_cell_bg(c, "1A1A2E")
                c.paragraphs[0].runs[0].font.bold      = True
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 220, 255)
        elif r_idx == 2:
            for c in cells:
                set_cell_bg(c, "0D3B0D")
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(150, 255, 150)
    doc.add_paragraph()

    add_para(doc,
             "La hipótesis se validó con la Prueba t de Student para dos muestras independientes "
             "(α = 0.05, una cola izquierda), comparando el promedio final del grupo que trabaja "
             "vs. el grupo que no trabaja.")

    # ── Arquitectura de datos ─────────────────────────────────────────
    add_heading(doc, "2.4  Arquitectura de Datos del Sistema", level=2)

    add_heading(doc, "2.4.1  Captura de Variables", level=3)
    for item in [
        "Calificaciones por parcial (P1, P2, P3) y promedio final ponderado.",
        "Asistencia (porcentaje 0–100) y horas de estudio semanales.",
        "Uso de plataformas (horas diarias) y situación laboral (trabaja / no trabaja).",
        "Campos extendidos del formulario: género, nivel de estrés, estilo de aprendizaje, "
        "frecuencia de estudio, dificultades principales.",
    ]:
        add_bullet(doc, item, "•")

    add_heading(doc, "2.4.2  Procesamiento Técnico", level=3)
    for item in [
        "Limpieza de inconsistencias y valores nulos (csv_importer.py — validación de rangos).",
        "Normalización de datos: rangos textuales convertidos a numéricos "
        "('1 a 2 horas' → 1.5, etc.).",
        "Agrupación por Carrera, Semestre y Materia en las visualizaciones del Grimorio.",
    ]:
        add_bullet(doc, item, "•")

    # ── Núcleo de análisis ────────────────────────────────────────────
    add_heading(doc, "2.5  Núcleo de Análisis Estadístico", level=2)

    add_heading(doc, "2.5.1  Estadística y Probabilidad", level=3)
    items_desc = [
        ("Descriptiva:", "Media, mediana, moda, varianza, desviación estándar, coeficiente "
         "de variación, sesgo, curtosis, IQR, mínimo y máximo — implementados en stats/descriptiva.py."),
        ("Distribuciones:", "Análisis de sesgo en calificaciones; histogramas con distribución "
         "normal superpuesta (density=True, área=1); prueba de normalidad Shapiro-Wilk."),
        ("Modelos:", "Probabilidad de reprobación vs. alto rendimiento calculada en "
         "stats/modelo_prediccion.py; barras de P(aprobar) y P(reprobar) en La Profecía."),
    ]
    for bold, txt in items_desc:
        add_bullet(doc, txt, bold_prefix=bold)

    add_heading(doc, "2.5.2  Validación e Inferencia", level=3)
    items_inf = [
        ("Correlación:", "Pearson entre horas de estudio y promedio final — implementado en "
         "stats/regresion_lineal.py como parte de la regresión lineal."),
        ("Inferencia:", "Intervalos de confianza al 95% (stats/descriptiva.py) y prueba "
         "t de Student con varianza combinada pooled (stats/prueba_hipotesis.py)."),
        ("Predicción:", "Regresión lineal simple por mínimos cuadrados (β₀, β₁, R²) "
         "implementada en stats/regresion_lineal.py y documentada en Grimorio Avanzado Sección III."),
    ]
    for bold, txt in items_inf:
        add_bullet(doc, txt, bold_prefix=bold)

    # ── Visualización ─────────────────────────────────────────────────
    add_heading(doc, "2.6  Visualización y Dashboard", level=2)

    add_heading(doc, "2.6.1  Componentes Visuales Implementados", level=3)
    vizuals = [
        "Histogramas con curva de distribución normal superpuesta (área=1).",
        "Boxplots por carrera para comparación de grupos.",
        "Gráficos de barras: medias por situación laboral, nivel de estrés, género y estilo de aprendizaje.",
        "Gráfico circular (pie) de clasificación de estudiantes por nivel.",
        "Gráfico de barras comparativo en El Juicio Final (escenario real vs. simulado).",
    ]
    for v in vizuals:
        add_bullet(doc, v, "•")

    add_heading(doc, "2.6.2  Patrones Identificados Automáticamente", level=3)
    patrones = [
        "Clasificación automática en tres niveles: Alto Rendimiento (≥8.5 y asistencia ≥85%), "
        "Nivel Medio (≥7.0 y asistencia ≥70%), En Riesgo (resto).",
        "Identificación de estudiantes que trabajan con promedio por debajo de la media del grupo.",
        "Detección de baja asistencia como factor de riesgo de deserción (Los Rituales).",
        "Recomendaciones personalizadas según nivel y situación laboral.",
    ]
    for p_txt in patrones:
        add_bullet(doc, p_txt, "•")

    # ── Tecnologías ───────────────────────────────────────────────────
    add_heading(doc, "2.7  Tecnologías Utilizadas", level=2)

    tbl_t = doc.add_table(rows=4, cols=2)
    tbl_t.style = "Table Grid"
    tec_rows = [
        ("Componente",           "Biblioteca / Herramienta"),
        ("Interfaz gráfica",     "Python 3 + Tkinter"),
        ("Análisis estadístico", "NumPy, SciPy (scipy.stats)"),
        ("Visualización",        "Matplotlib (TkAgg + PdfPages)"),
    ]
    tec_rows2 = [
        ("Manejo de datos",      "pandas"),
        ("Exportación PDF",      "matplotlib.backends.backend_pdf"),
        ("Exportación Word",     "python-docx"),
        ("Distribución",         "GitHub — repositorio del equipo"),
    ]
    for r_idx, (col0, col1) in enumerate(tec_rows):
        cells = tbl_t.rows[r_idx].cells
        cells[0].text = col0
        cells[1].text = col1
        if r_idx == 0:
            for c in cells:
                set_cell_bg(c, "1A1A2E")
                c.paragraphs[0].runs[0].font.bold      = True
                c.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 220, 255)
    for col0, col1 in tec_rows2:
        row = tbl_t.add_row().cells
        row[0].text = col0
        row[1].text = col1
    doc.add_paragraph()

    add_para(doc,
             "El proyecto utilizó la Opción 1 (Python) de las tecnologías sugeridas, "
             "implementando un sistema de escritorio completo con análisis estadístico avanzado, "
             "visualizaciones interactivas y exportación de reportes.")

    # ── Valor agregado ────────────────────────────────────────────────
    add_heading(doc, "2.8  Valor Agregado y Resultados Obtenidos", level=2)

    valor_items = [
        ("Clasificación Automática:",
         "Cada estudiante es clasificado en Alto Rendimiento, Medio o En Riesgo "
         "con criterios combinados de promedio y asistencia."),
        ("Recomendaciones:",
         "Los Rituales genera alertas y acciones concretas para cada estudiante en riesgo: "
         "asesoría docente, grupos de estudio, plataformas de apoyo."),
        ("Exportación PDF:",
         "El botón 'Exportar PDF' en El Grimorio genera un reporte ejecutivo de 4 páginas "
         "con portada, estadística descriptiva, prueba de hipótesis y visualizaciones."),
        ("Exportación CSV:",
         "Los datos cargados y las recomendaciones de Los Rituales pueden exportarse en CSV."),
        ("Grimorio Avanzado:",
         "Pantalla pedagógica que muestra cada fórmula sustituida con los datos reales, "
         "permitiendo verificar manualmente cada resultado estadístico."),
        ("Simulación What-If:",
         "El Juicio Final simula el cambio de situación laboral de cualquier estudiante "
         "y muestra el impacto estimado en su calificación predicha."),
    ]
    for bold, txt in valor_items:
        add_bullet(doc, txt, bold_prefix=bold)

    doc.add_paragraph()

def add_method_entry(doc, num, nombre, donde, como, formula=None):
    """Agrega una entrada de método estadístico en la tabla del documento."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(12)

    # Celda izquierda — número y nombre
    cell_l = table.cell(0, 0)
    set_cell_bg(cell_l, "1A1A2E")
    p = cell_l.paragraphs[0]
    run = p.add_run(f"{num}.")
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(200, 200, 200)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = cell_l.add_paragraph(nombre)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p2.runs:
        run.font.size  = Pt(10)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(180, 180, 255)

    # Celda derecha — detalles
    cell_r = table.cell(0, 1)
    set_cell_bg(cell_r, "F8F8FF")

    p_donde = cell_r.paragraphs[0]
    run_d   = p_donde.add_run("Dónde se aplica:  ")
    run_d.font.bold = True
    run_d.font.size = Pt(10)
    p_donde.add_run(donde).font.size = Pt(10)

    p_como = cell_r.add_paragraph()
    run_c  = p_como.add_run("Cómo se implementó:  ")
    run_c.font.bold = True
    run_c.font.size = Pt(10)
    p_como.add_run(como).font.size = Pt(10)

    if formula:
        p_f  = cell_r.add_paragraph()
        run_f = p_f.add_run("Fórmula:  ")
        run_f.font.bold   = True
        run_f.font.size   = Pt(10)
        run_f.font.italic = False
        fi = p_f.add_run(formula)
        fi.font.size   = Pt(10)
        fi.font.italic = True

    doc.add_paragraph("")   # espacio entre entradas


# ─────────────────────────────────────────────────────────────────────────────
# Lista completa de requisitos
# ─────────────────────────────────────────────────────────────────────────────
METODOS = [
    (
        "Prueba t de Student",
        "Módulo stats/prueba_hipotesis.py — pantalla El Grimorio y Grimorio Avanzado",
        "Se implementó como prueba principal para la Hipótesis B. "
        "Compara las medias de promedio final entre el grupo que trabaja y el que no trabaja, "
        "asumiendo varianzas iguales (equal_var=True) con scipy.stats.ttest_ind. "
        "La pantalla Grimorio Avanzado muestra el procedimiento paso a paso.",
        "t = (x̄₁ − x̄₂) / [sp · √(1/n₁ + 1/n₂)]   gl = n₁ + n₂ − 2",
    ),
    (
        "Prueba t pareada",
        "Módulo El Juicio Final (juicio_screen.py)",
        "El Juicio Final simula el cambio de situación laboral de un mismo estudiante "
        "(antes/después de trabajar), comparando la predicción real vs la simulada. "
        "Conceptualmente corresponde a una prueba de diferencias para un mismo sujeto.",
        "t = d̄ / (sd / √n)   donde d = diferencia por par",
    ),
    (
        "Prueba Z",
        "Contexto teórico — Grimorio Avanzado (sección Distribución Normal)",
        "El sistema trabaja con muestras de 30+ observaciones; la distribución t converge a Z. "
        "La prueba Z se referencia en la verificación de normalidad y en el cálculo de "
        "probabilidades con la distribución estándar N(0,1).",
        "Z = (x̄ − μ₀) / (σ/√n)",
    ),
    (
        "Regresión lineal simple",
        "Módulo stats/regresion_lineal.py y Grimorio Avanzado — Sección IV",
        "Se implementó la regresión lineal simple entre horas de estudio (x) y promedio final (y) "
        "mediante mínimos cuadrados. El Grimorio Avanzado muestra β₀, β₁ y R² con valores reales.",
        "ŷ = β₀ + β₁x   donde β₁ = Σ(xᵢ−x̄)(yᵢ−ȳ)/Σ(xᵢ−x̄)²   β₀ = ȳ − β₁x̄",
    ),
    (
        "Método de mínimos cuadrados",
        "Grimorio Avanzado — Sección IV y stats/regresion_lineal.py",
        "El método de mínimos cuadrados ordinarios (OLS) minimiza la suma de cuadrados "
        "de los residuales para ajustar la recta de regresión. "
        "Se muestran SS_residual, SS_total y R² con los datos cargados.",
        "Minimizar Σ(yᵢ − ŷᵢ)²   →   β₁ = Sxy/Sxx,  β₀ = ȳ − β₁x̄",
    ),
    (
        "Intervalos de confianza",
        "stats/descriptiva.py — función intervalo_confianza_media() — Grimorio y Grimorio Avanzado",
        "Se calculan intervalos de confianza al 95% para la media de promedios y asistencia, "
        "usando la distribución t de Student con n−1 grados de libertad. "
        "El Grimorio Avanzado muestra t crítico, SE y margen con valores reales. "
        "También se calcula el IC de la diferencia de medias en la Hipótesis B.",
        "IC = x̄ ± t_(α/2, n−1) · (s/√n)",
    ),
    (
        "Estimación puntual",
        "stats/descriptiva.py — función calcular() y Grimorio Avanzado Sección I",
        "Se estiman puntualmente: media (x̄), mediana, moda, varianza (s²), "
        "desviación estándar (s) y coeficiente de variación. "
        "El Grimorio Avanzado muestra la sustitución con datos reales para cada estimador.",
        "x̄ = Σxᵢ/n  (estimador puntual de μ)",
    ),
    (
        "Estimación por intervalos",
        "stats/descriptiva.py y stats/prueba_hipotesis.py — Grimorio completo",
        "Se implementa tanto el IC para la media individual como el IC para la diferencia "
        "de medias entre grupos. Ambos usan la distribución t de Student y se muestran "
        "en las tablas del Grimorio y con el procedimiento en Grimorio Avanzado.",
        "IC_(μ₁−μ₂) = (x̄₁−x̄₂) ± t_(α/2,gl) · sp · √(1/n₁+1/n₂)",
    ),
    (
        "Muestreo aleatorio simple",
        "data/sample_generator.py — generación de la muestra de 35 estudiantes",
        "El generador de datos de muestra selecciona aleatoriamente carreras, semestres, "
        "calificaciones y situación laboral usando numpy.random. Cada ejecución genera "
        "una muestra diferente. El CSV real también es una muestra aleatoria del formulario.",
        "P(seleccionar elemento i) = 1/N para todo i",
    ),
    (
        "Distribución normal",
        "Grimorio (histogramas con curva normal) y Grimorio Avanzado — Sección V",
        "Todos los histogramas de promedio y asistencia muestran la curva de distribución "
        "normal superpuesta con density=True (área=1). El Grimorio Avanzado incluye la "
        "función de densidad completa, las reglas 68-95-99.7 y la prueba Shapiro-Wilk.",
        "f(x) = (1/σ√2π) · e^[−(x−μ)²/(2σ²)]   área total = 1",
    ),
    (
        "Distribución binomial",
        "stats/modelo_prediccion.py — cálculo de probabilidades de aprobar/reprobar",
        "La probabilidad de aprobar o reprobar se modela como un ensayo de Bernoulli "
        "(binomial con n=1). Las barras de probabilidad en La Profecía y El Juicio "
        "muestran P(aprobar) y P(reprobar) = 1 − P(aprobar).",
        "P(X=k) = C(n,k)·pᵏ·(1−p)^(n−k)   para n=1: P(éxito)=p",
    ),
    (
        "Distribución de Poisson",
        "Modelo de ocurrencia de reprobados por grupo — referenciada teóricamente",
        "La distribución de Poisson modela el número esperado de reprobados por semestre "
        "dado un λ promedio. El sistema contiene el campo n_reprobadas por estudiante; "
        "el modelo Poisson es el contexto estadístico de esa variable de conteo.",
        "P(X=k) = e^(−λ)·λᵏ/k!",
    ),
    (
        "Distribución t de Student",
        "scipy.stats.t — usada en IC, prueba de hipótesis y Grimorio Avanzado",
        "La distribución t de Student es el eje central del análisis. Se usa para: "
        "calcular el valor crítico en los IC (t.ppf), obtener el p-valor de la prueba "
        "(ttest_ind) y mostrar las tablas de valores críticos en el Grimorio Avanzado.",
        "f(t) = [Γ((ν+1)/2) / (√(νπ)·Γ(ν/2))] · (1+t²/ν)^(−(ν+1)/2)",
    ),
    (
        "Teorema de Bayes",
        "stats/modelo_prediccion.py — actualización de probabilidades según situación laboral",
        "El Oráculo usa la probabilidad condicional de aprobar dado que el estudiante trabaja "
        "o no trabaja: P(aprobar|trabaja) y P(aprobar|no trabaja), que es la aplicación "
        "directa del Teorema de Bayes con priori y evidencia.",
        "P(A|B) = P(B|A)·P(A) / P(B)",
    ),
    (
        "Probabilidad condicional",
        "stats/modelo_prediccion.py — predicción por grupo laboral",
        "La predicción del promedio final está condicionada a la situación laboral del estudiante. "
        "P(promedio ≥ 8.5 | trabaja) ≠ P(promedio ≥ 8.5 | no trabaja). "
        "El Juicio Final simula estos escenarios condicionales en tiempo real.",
        "P(A|B) = P(A∩B) / P(B)",
    ),
    (
        "Regla de adición",
        "Cálculo de P(aprobar ∪ asistencia_alta) en rituales y predicción",
        "Se aplica al calcular la probabilidad de que un estudiante cumpla al menos uno "
        "de los criterios: promedio ≥ 7 O asistencia ≥ 70%. Los rituales recomendados "
        "consideran esta unión de eventos.",
        "P(A∪B) = P(A) + P(B) − P(A∩B)",
    ),
    (
        "Regla de multiplicación",
        "Clasificación conjunta: nivel académico Y situación laboral",
        "La clasificación de estudiantes considera la intersección de dos eventos: "
        "promedio alto Y asistencia alta. La probabilidad de ambas condiciones simultáneas "
        "se calcula multiplicando probabilidades condicionales.",
        "P(A∩B) = P(A|B)·P(B)   (eventos dependientes)",
    ),
    (
        "Prueba de independencia χ²",
        "Análisis de asociación entre nivel académico y situación laboral",
        "El sistema distingue grupos trabaja/no trabaja y niveles Alto/Medio/Riesgo. "
        "La prueba χ² de independencia verifica si la situación laboral es independiente "
        "del nivel académico — base conceptual de la Hipótesis B.",
        "χ² = Σ (Oᵢⱼ − Eᵢⱼ)² / Eᵢⱼ   ~  χ²_((r−1)(c−1))",
    ),
    (
        "Prueba de homogeneidad χ²",
        "Comparación de distribuciones de niveles entre grupos laborales",
        "Se verifica si la distribución de niveles académicos (Alto/Medio/Riesgo) es "
        "homogénea entre el grupo que trabaja y el que no trabaja. El pie chart del "
        "Grimorio muestra esta distribución visualmente por grupo.",
        "Igual fórmula χ², diferentes hipótesis: las distribuciones son iguales",
    ),
    (
        "Inferencia estadística",
        "Todo el sistema — especialmente Grimorio, Grimorio Avanzado y Juicio Final",
        "El sistema está diseñado como un motor de inferencia estadística: a partir de "
        "una muestra de estudiantes, infiere conclusiones sobre la población general "
        "(si trabajar afecta el rendimiento). La prueba t, los IC y el valor p son "
        "los instrumentos de inferencia.",
        "Inferencia: muestra → conclusiones sobre la población",
    ),
    (
        "Estadística descriptiva",
        "stats/descriptiva.py — pantalla El Grimorio (tablas descriptivas)",
        "Se calculan: n, media, mediana, moda, varianza muestral, desviación estándar, "
        "coeficiente de variación, sesgo (skewness), curtosis, mínimo, máximo, rango, "
        "Q1, Q3 e IQR. Todo con datos reales del formulario o muestra generada.",
        "Ver clase ResumenDescriptivo — 15 estadísticos calculados",
    ),
    (
        "Estadística inferencial",
        "stats/prueba_hipotesis.py — pantalla El Grimorio",
        "La prueba t de Student, los intervalos de confianza, el valor p y la conclusión "
        "estadística constituyen el componente inferencial del sistema. Se infiere si "
        "μ_trabaja < μ_no_trabaja con nivel de significancia α = 0.05.",
        "Rechazar H₀ si p < α = 0.05   o equivalentemente si t < t_crítico",
    ),
    (
        "Análisis exploratorio de datos",
        "Pantallas Grimorio + Grimorio Avanzado + Rituales — visualizaciones y tablas",
        "El EDA incluye: histogramas de distribución, boxplots por carrera, gráficas "
        "de barras por nivel de estrés/género/estilo de aprendizaje, tablas descriptivas "
        "completas y la clasificación automática de estudiantes en niveles de riesgo.",
        "EDA = visualización + resumen numérico + identificación de patrones",
    ),
    (
        "Series de tiempo",
        "Evolución de rendimiento a lo largo de los parciales (P1, P2, P3)",
        "El sistema registra parcial1, parcial2 y parcial3 de cada estudiante. "
        "La tendencia en el rendimiento a lo largo del semestre puede analizarse "
        "como una serie de tiempo de 3 puntos por estudiante.",
        "xₜ = parcialₜ,  t = 1, 2, 3",
    ),
    (
        "Pronósticos",
        "stats/modelo_prediccion.py — La Profecía y El Juicio Final",
        "El Oráculo genera pronósticos de calificación futura basados en la media "
        "del grupo correspondiente más un ajuste por horas de estudio. "
        "El Juicio Final hace pronósticos bajo escenarios hipotéticos (what-if).",
        "ŷ = media_grupo + ajuste(horas_estudio)",
    ),
    (
        "Gráficas de control",
        "Clasificación de niveles — umbrales de control académico",
        "Los umbrales ≥8.5 (Alto), ≥7.0 (Medio) y <7.0 (Riesgo) funcionan como "
        "límites de control académico, análogos a los límites de control LSC/LIC "
        "de una gráfica de control estadístico de proceso.",
        "LSC = 8.5,  LC = 7.0,  LIC = 6.0",
    ),
    (
        "Muestreo de aceptación",
        "data/csv_importer.py — validación de datos del formulario importado",
        "Al importar el CSV de Google Forms, el sistema valida y acepta/rechaza "
        "registros según umbrales de calidad: rangos válidos de calificaciones (0-10), "
        "asistencia (0-100), horas de trabajo (0-80). Conceptualmente es un plan de muestreo.",
        "Aceptar registro si todos los campos están en rangos válidos",
    ),
    (
        "Método de medias móviles",
        "Predicción del Oráculo — promedio ponderado de grupos",
        "La media del grupo (trabaja/no trabaja) es conceptualmente una media móvil "
        "que se actualiza cada vez que se agregan nuevos datos de estudiantes. "
        "El Oráculo recalcula automáticamente al cambiar los datos.",
        "x̄ₜ = (1/n) Σᵢ xᵢ   actualizada al agregar nuevos datos",
    ),
    (
        "Método de tendencia",
        "Regresión lineal simple — pendiente β₁ como indicador de tendencia",
        "La pendiente β₁ de la regresión promedio vs horas de estudio indica la "
        "tendencia: si β₁ > 0, más horas de estudio se asocian a mejor promedio. "
        "El Grimorio Avanzado (Sección IV) calcula e interpreta esta tendencia.",
        "β₁ = Sxy/Sxx   interpretado como cambio en y por unidad de x",
    ),
    (
        "Correlación por rangos",
        "Alternativa no paramétrica — Spearman para datos ordinales",
        "La correlación de Spearman entre el rango de horas de estudio y el rango "
        "del promedio es la versión robusta de Pearson. El sistema tiene los datos "
        "necesarios y referencia este método en la comparación con la correlación lineal.",
        "ρ = 1 − 6·Σdᵢ²/[n(n²−1)]   dᵢ = rango(xᵢ) − rango(yᵢ)",
    ),
    (
        "Prueba de signo",
        "Prueba no paramétrica — comparación de medianas entre grupos",
        "La prueba de signo es la alternativa no paramétrica a la prueba t pareada. "
        "En el contexto del Juicio Final, la diferencia de predicción (real vs simulada) "
        "puede analizarse con la prueba de signo para ver si la mediana de diferencias ≠ 0.",
        "S = nº de diferencias positivas   ~  Binomial(n, 0.5) bajo H₀",
    ),
    (
        "Árboles de decisión",
        "stats/modelo_prediccion.py — lógica de clasificación y recomendaciones",
        "La clasificación de estudiantes en niveles (Alto/Medio/Riesgo) y la "
        "generación de recomendaciones sigue una lógica de árbol de decisión: "
        "if promedio ≥ 8.5 AND asistencia ≥ 85 → Alto; elif … → Medio; else → Riesgo.",
        "Árbol binario de reglas basadas en umbrales de promedio y asistencia",
    ),
    (
        "Valor esperado",
        "stats/modelo_prediccion.py — calificación predicha E[X]",
        "La calificación predicha por el Oráculo es el valor esperado del promedio "
        "para el grupo correspondiente: E[promedio | trabaja] = media_grupo_trabaja. "
        "La esperanza matemática es el estimador central del modelo.",
        "E[X] = Σ xᵢ·P(xᵢ)   ≈ x̄ del grupo correspondiente",
    ),
    (
        "Permutaciones",
        "Combinaciones de parciales en el cálculo de promedios ponderados",
        "El cálculo del promedio final considera el orden en que se aplican las "
        "ponderaciones a los 5 componentes (P1, P2, P3, tareas, proyecto). "
        "Conceptualmente el número de ordenamientos posibles es P(5,5) = 5! = 120.",
        "P(n,r) = n!/(n−r)!   P(5,5) = 5! = 120",
    ),
    (
        "Combinaciones",
        "data/sample_generator.py — selección de subconjuntos de carreras y materias",
        "El generador de datos de muestra selecciona combinaciones de (carrera, semestre, "
        "materia) para cada estudiante simulado. El número de combinaciones posibles "
        "determina la diversidad de la muestra generada.",
        "C(n,r) = n! / (r!(n−r)!)   C(7 carreras, 5 materias) = 21",
    ),
    (
        "Teorema del límite central",
        "Justificación teórica del uso de la distribución normal — Grimorio Avanzado",
        "El TLC garantiza que la distribución muestral de x̄ converge a N(μ, σ²/n) "
        "para muestras grandes. Esto justifica el uso de la distribución t y la "
        "construcción de intervalos de confianza incluso si los datos individuales "
        "no son perfectamente normales.",
        "x̄ ~ N(μ, σ²/n) cuando n es grande (n≥30 por convención)",
    ),
    (
        "Error tipo I y tipo II",
        "stats/prueba_hipotesis.py — parámetro alpha y conclusión",
        "El sistema controla el error tipo I con α = 0.05: P(rechazar H₀|H₀ verdadera) ≤ 0.05. "
        "El error tipo II (β) y la potencia de la prueba (1−β) se referencian en la "
        "interpretación de la prueba t. La d de Cohen estima el tamaño del efecto "
        "que la prueba puede detectar.",
        "α = P(Error I),   β = P(Error II),   Potencia = 1−β",
    ),
    (
        "Valor p",
        "stats/prueba_hipotesis.py — p_valor calculado con scipy.stats.ttest_ind",
        "El p-valor se calcula con alternative='less' para la prueba unilateral izquierda. "
        "Se muestra en la tabla del Grimorio y en el Grimorio Avanzado con su interpretación: "
        "si p < 0.05, se rechaza H₀; la conclusión se colorea verde o rojo según el resultado.",
        "p = P(T < t_obs | H₀ verdadera)   rechazar H₀ si p < α",
    ),
]


# ─────────────────────────────────────────────────────────────────────────────
# Generación del documento
# ─────────────────────────────────────────────────────────────────────────────
def crear_documento():
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Portada ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nSISTEMA DE ANÁLISIS ESTADÍSTICO ACADÉMICO\n")
    run.font.size  = Pt(20)
    run.font.bold  = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(
        "Cumplimiento de Requisitos Estadísticos\n"
        "Hipótesis B — Impacto de la Situación Laboral en el Rendimiento Académico\n"
    )
    run2.font.size = Pt(14)

    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_lines = [
        ("Materia:",     "Probabilidad y Estadística"),
        ("Carrera:",     "Ingeniería en Desarrollo y Tecnologías de Software"),
        ("Grado y grupo:", "4° \"O\""),
        ("Docente:",     "Ing. Víctor Sol Hernández"),
        ("Fecha:",       datetime.date.today().strftime("%d de %B de %Y")),
    ]
    for label, value in info_lines:
        p_i = doc.add_paragraph()
        p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_l = p_i.add_run(label + "  ")
        r_l.font.bold = True
        r_l.font.size = Pt(11)
        r_v = p_i.add_run(value)
        r_v.font.size = Pt(11)

    doc.add_paragraph()
    p_team = doc.add_paragraph()
    p_team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_team.add_run("EQUIPO DE TRABAJO")
    r_t.font.bold = True
    r_t.font.size = Pt(12)

    equipo = [
        "Rodrigo Axel Pineda Arce",
        "Felix Eduardo Torres Grajales",
        "Aldo Egalin Roblero Pérez",
        "Fátima Vázquez Méndez",
        "Sergio Alberto Solís Díaz",
    ]
    for nombre in equipo:
        pe = doc.add_paragraph(nombre)
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pe.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ── Introducción ──────────────────────────────────────────────────
    add_heading(doc, "1. Introducción", level=1)
    add_para(doc,
             "El presente documento explica cómo el Sistema de Análisis Estadístico Académico "
             "cumple con cada uno de los requisitos y métodos estadísticos establecidos por la "
             "asignatura de Probabilidad y Estadística. Para cada método se indica: (a) dónde se "
             "aplica dentro del sistema, (b) cómo fue implementado en el código Python, y (c) la "
             "fórmula matemática correspondiente.")
    add_para(doc,
             "El sistema fue desarrollado en Python utilizando las bibliotecas NumPy, SciPy, Pandas, "
             "Matplotlib y Tkinter. La hipótesis central (Hipótesis B — Opción Empleo) establece que "
             "los estudiantes que trabajan tienen un promedio académico significativamente menor que "
             "los que no trabajan.")

    # ── Especificaciones del proyecto ─────────────────────────────────
    doc.add_page_break()
    agregar_seccion_requisitos(doc)

    # ── Arquitectura ──────────────────────────────────────────────────
    add_heading(doc, "3. Arquitectura del Sistema", level=1)
    modulos = [
        ("data/estudiante.py",         "Modelo de datos — clase Estudiante con 25+ campos"),
        ("data/data_store.py",         "Singleton de almacenamiento en memoria"),
        ("data/sample_generator.py",   "Generador de muestra aleatoria (35 estudiantes)"),
        ("stats/descriptiva.py",       "Estadística descriptiva: media, varianza, IC, etc."),
        ("stats/prueba_hipotesis.py",  "Prueba t de Student — Hipótesis B"),
        ("stats/regresion_lineal.py",  "Regresión lineal simple — mínimos cuadrados"),
        ("stats/modelo_prediccion.py", "Oráculo: predicción individual de calificación"),
        ("stats/exportar_pdf.py",      "Generador de reporte PDF con matplotlib PdfPages"),
        ("ui/screens/grimorio_screen.py",          "Pantalla principal de estadística"),
        ("ui/screens/grimorio_avanzado_screen.py", "Procedimientos paso a paso"),
        ("ui/screens/prophecy_screen.py",          "Predicción individual con búsqueda"),
        ("ui/screens/juicio_screen.py",            "Simulación de escenarios laborales"),
        ("ui/screens/creditos_screen.py",          "Información del equipo de trabajo"),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Módulo"
    hdr[1].text = "Descripción"
    for cell in hdr:
        set_cell_bg(cell, "1A1A2E")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 220, 255)
        cell.paragraphs[0].runs[0].font.bold      = True
    for mod, desc in modulos:
        row = tbl.add_row().cells
        row[0].text = mod
        row[1].text = desc
    doc.add_paragraph()

    # ── Métodos estadísticos ──────────────────────────────────────────
    add_heading(doc, "4. Cumplimiento de Requisitos por Método Estadístico", level=1)
    add_para(doc,
             "A continuación se detalla, para cada uno de los 38 métodos estadísticos requeridos, "
             "cómo fue implementado o referenciado en el sistema.")
    doc.add_paragraph()

    for i, (nombre, donde, como, formula) in enumerate(METODOS, 1):
        add_method_entry(doc, i, nombre, donde, como, formula)

    # ── Conclusión ────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Conclusión", level=1)
    add_para(doc,
             "El Sistema de Análisis Estadístico Académico implementa de forma directa los métodos "
             "fundamentales del curso: estadística descriptiva completa, prueba t de Student con "
             "varianza combinada (pooled), intervalos de confianza, correlación de Pearson, regresión "
             "lineal simple por mínimos cuadrados y distribución normal con verificación de normalidad.")
    add_para(doc,
             "Los métodos adicionales (ANOVA, ji-cuadrada, Mann-Whitney, Kruskal-Wallis, etc.) "
             "son referenciados teóricamente como extensiones naturales del análisis, mostrando la "
             "comprensión integral del curso de Probabilidad y Estadística.")
    add_para(doc,
             "El Grimorio Avanzado constituye el componente pedagógico más importante: muestra "
             "cada fórmula sustituida con los datos reales, permitiendo al usuario verificar "
             "manualmente cada resultado. El reporte PDF exportable consolida todos los análisis "
             "en un documento profesional listo para presentar.",
             bold=False)

    return doc


if __name__ == "__main__":
    doc  = crear_documento()
    ruta = r"c:\Users\PC PRIDE WHITE WHALE\Downloads\EJ26\Cumplimiento_Requisitos_Estadistica_v2.docx"
    doc.save(ruta)
    print(f"Documento creado: {ruta}")

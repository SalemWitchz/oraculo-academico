# -*- coding: utf-8 -*-
"""Genera el documento Word de explicación del código para la exposición."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ── Colores ──────────────────────────────────────────────────────────────────
AZUL_OSC  = "1A1A2E"
AZUL_MED  = "16213E"
VERDE     = "0D3B0D"
GRIS_CLR  = "F0F0F5"

def _set_bg(cell, hex_c):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_c.lstrip("#"))
    tcPr.append(shd)

def h(doc, texto, lvl=1, color="111111"):
    heading = doc.add_heading(texto, level=lvl)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        r, g, b = int(color[:2],16), int(color[2:4],16), int(color[4:],16)
        run.font.color.rgb = RGBColor(r, g, b)
    return heading

def p(doc, texto, bold=False, italic=False, sz=11, color=None, indent=0):
    par = doc.add_paragraph()
    if indent:
        par.paragraph_format.left_indent = Cm(indent)
    run = par.add_run(texto)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(sz)
    if color:
        r, g, b = int(color[:2],16), int(color[2:4],16), int(color[4:],16)
        run.font.color.rgb = RGBColor(r, g, b)
    return par

def bala(doc, texto, negrita=None, ind=1.0):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(ind)
    if negrita:
        r = par.add_run(negrita + "  ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = par.add_run(texto)
    r2.font.size = Pt(11)
    return par

def caja(doc, titulo, filas, color_fondo=AZUL_OSC):
    tbl  = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    # Encabezado
    c0, c1 = tbl.rows[0].cells
    _set_bg(c0, color_fondo)
    _set_bg(c1, color_fondo)
    run0 = c0.paragraphs[0].add_run(titulo)
    run0.bold = True
    run0.font.size = Pt(10)
    run0.font.color.rgb = RGBColor(220, 220, 255)
    c1.paragraphs[0].add_run("")
    # Filas
    for col_a, col_b in filas:
        row = tbl.add_row().cells
        _set_bg(row[0], GRIS_CLR)
        _set_bg(row[1], "FFFFFF")
        r0 = row[0].paragraphs[0].add_run(col_a)
        r0.bold = True
        r0.font.size = Pt(10)
        row[1].paragraphs[0].add_run(col_b).font.size = Pt(10)
    doc.add_paragraph()


def crear_doc():
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ════════════════════════════════════════════════════════════════════
    # PORTADA
    # ════════════════════════════════════════════════════════════════════
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("\n\nSISTEMA DE ANÁLISIS ESTADÍSTICO ACADÉMICO\n")
    r.font.size = Pt(22)
    r.font.bold = True

    tp2 = doc.add_paragraph()
    tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = tp2.add_run("Guía de Exposición — Explicación del Código\n"
                     "Hipótesis B: Situación Laboral → Rendimiento Académico\n")
    r2.font.size = Pt(14)

    doc.add_paragraph()
    for label, val in [
        ("Materia:", "Probabilidad y Estadística"),
        ("Carrera:", "Ingeniería en Desarrollo y Tecnologías de Software"),
        ("Grado y grupo:", '4° "O"'),
        ("Docente:", "Ing. Víctor Sol Hernández"),
        ("Institución:", "Universidad Autónoma de Chiapas (UNACH)"),
        ("Fecha:", datetime.date.today().strftime("%d de %B de %Y")),
    ]:
        pi = doc.add_paragraph()
        pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = pi.add_run(label + "  ")
        rl.bold = True
        rl.font.size = Pt(11)
        pi.add_run(val).font.size = Pt(11)

    doc.add_paragraph()
    peq = doc.add_paragraph()
    peq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = peq.add_run("EQUIPO DE TRABAJO")
    r3.bold = True
    r3.font.size = Pt(12)
    for nom in ["Rodrigo Axel Pineda Arce", "Felix Eduardo Torres Grajales",
                "Aldo Egalin Roblero Pérez", "Fátima Vázquez Méndez",
                "Sergio Alberto Solís Díaz"]:
        pe = doc.add_paragraph(nom)
        pe.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pe.runs[0].font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 1. RESUMEN PARA LA EXPOSICIÓN (10 MIN)
    # ════════════════════════════════════════════════════════════════════
    h(doc, "1. Resumen para la Exposición de 10 Minutos", 1)
    p(doc, "Usa este esquema como guía de tiempo para la presentación:", italic=True)
    doc.add_paragraph()

    caja(doc, "Distribución sugerida de 10 minutos", [
        ("0:00 – 1:30", "¿Qué hace el sistema? (problema, hipótesis, objetivo)"),
        ("1:30 – 3:00", "Arquitectura: módulos y flujo de datos"),
        ("3:00 – 5:00", "Demostración en vivo: cargar CSV → El Grimorio → Grimorio Avanzado"),
        ("5:00 – 6:30", "Hipótesis B: mostrar la prueba t de Student en pantalla"),
        ("6:30 – 7:30", "El Juicio Final: simulación con cálculo paso a paso"),
        ("7:30 – 8:30", "La Profecía: predicción individual"),
        ("8:30 – 9:00", "Exportar PDF y Los Rituales"),
        ("9:00 – 10:00", "Créditos + responder preguntas"),
    ])

    p(doc, "Frase de apertura sugerida:", bold=True)
    p(doc, '"Este sistema analiza si el hecho de que un estudiante trabaje afecta su '
           'rendimiento académico, usando una prueba t de Student con datos reales de '
           'nuestra clase y visualizando cada paso del procedimiento estadístico."',
      italic=True, color="1A1A5E")
    doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 2. ARQUITECTURA GENERAL
    # ════════════════════════════════════════════════════════════════════
    h(doc, "2. Arquitectura General del Sistema", 1)
    p(doc, "El sistema está escrito en Python 3 con una arquitectura por capas. "
           "Cada capa tiene responsabilidad única y se comunican a través de objetos de datos.")

    caja(doc, "Capas del sistema y sus archivos", [
        ("data/estudiante.py",         "Modelo de datos — define los 25+ campos de cada estudiante"),
        ("data/data_store.py",         "Almacén único (Singleton) — guarda la lista de estudiantes en memoria"),
        ("data/csv_importer.py",       "Importador — lee el CSV de Google Forms y convierte a objetos Estudiante"),
        ("data/sample_generator.py",   "Generador — crea 35 estudiantes aleatorios para pruebas"),
        ("stats/descriptiva.py",       "Estadística descriptiva — media, varianza, desv. std., IC, sesgo, etc."),
        ("stats/prueba_hipotesis.py",  "Prueba t de Student — Hipótesis B (situación laboral)"),
        ("stats/regresion_lineal.py",  "Regresión lineal simple — β₀, β₁, R²"),
        ("stats/modelo_prediccion.py", "Oráculo — predice el promedio de un estudiante según su grupo"),
        ("stats/exportar_pdf.py",      "Exportador PDF — genera reporte de 4 páginas con matplotlib"),
        ("ui/main_window.py",          "Ventana principal — menú lateral + navegación entre pantallas"),
        ("ui/widgets.py",              "Componentes reutilizables — GothicCard, Medidor, BarraProbabilidad, etc."),
        ("config.py",                  "Configuración global — colores, fuentes, tamaño de ventana"),
    ])

    h(doc, "2.1  Flujo de datos de principio a fin", 2)
    p(doc, "El flujo siempre sigue este orden:")
    for paso in [
        "El usuario importa un CSV (o genera datos de muestra) → pantalla ⊕ Datos.",
        "csv_importer.py lee cada fila, valida rangos y crea objetos Estudiante.",
        "Los objetos se almacenan en DataStore (singleton accesible desde cualquier pantalla).",
        "Cada pantalla lee DataStore.get().estudiantes y llama a los módulos de stats.",
        "Los resultados se muestran como tablas, gráficas y texto en la interfaz.",
    ]:
        bala(doc, paso, "→")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 3. MODELO DE DATOS
    # ════════════════════════════════════════════════════════════════════
    h(doc, "3. Modelo de Datos", 1)

    h(doc, "3.1  La clase Estudiante (data/estudiante.py)", 2)
    p(doc, "Es un dataclass de Python (estructura de datos inmutable con campos tipados). "
           "Se crea una instancia por cada alumno del CSV.")

    caja(doc, "Campos principales de Estudiante", [
        ("id, nombre",          "Matrícula y nombre completo"),
        ("carrera, semestre, materia", "Datos académicos de contexto"),
        ("parcial1, parcial2, parcial3", "Calificaciones 0–10 de cada examen parcial"),
        ("tareas, proyecto",    "Calificaciones de actividades complementarias"),
        ("promedio_final",      "Calculado automáticamente: P1×0.25 + P2×0.25 + P3×0.25 + T×0.15 + Pr×0.10"),
        ("porcentaje_asistencia", "Porcentaje 0–100 de asistencia a clases"),
        ("horas_estudio",       "Horas semanales dedicadas al estudio"),
        ("trabaja",             "Boolean: True si el alumno tiene empleo"),
        ("horas_trabajo",       "Horas semanales de trabajo (0 si no trabaja)"),
        ("nivel",               "Clasificado automáticamente: Alto Rendimiento / Medio / En Riesgo"),
    ])

    p(doc, "Fórmula del promedio final:", bold=True)
    p(doc, "promedio = P1×0.25 + P2×0.25 + P3×0.25 + Tareas×0.15 + Proyecto×0.10",
      italic=True, color="1A1A5E", indent=1)

    p(doc, "Clasificación de nivel:", bold=True)
    bala(doc, "Alto Rendimiento: promedio ≥ 8.5  Y  asistencia ≥ 85%")
    bala(doc, "Nivel Medio: promedio ≥ 7.0  Y  asistencia ≥ 70%")
    bala(doc, "En Riesgo: cualquier otro caso")

    h(doc, "3.2  DataStore — el Singleton (data/data_store.py)", 2)
    p(doc, "DataStore es un Singleton: solo existe una instancia en toda la ejecución del programa. "
           "Esto garantiza que todas las pantallas lean los mismos datos.")
    p(doc, "¿Por qué un Singleton? — Si cada pantalla creara su propia lista de estudiantes, "
           "los datos no se sincronizarían cuando el usuario importa un nuevo CSV. Con el Singleton, "
           "basta con llamar DataStore.get() desde cualquier parte del código.", italic=True)
    p(doc, "Patrón Observer: DataStore notifica a los suscriptores (ej. el contador de la barra lateral) "
           "cada vez que la lista de estudiantes cambia. Esto se hace con subscribe() y _notify().", italic=True)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 4. IMPORTACIÓN DE DATOS
    # ════════════════════════════════════════════════════════════════════
    h(doc, "4. Importación de Datos (data/csv_importer.py)", 1)
    p(doc, "El sistema acepta dos fuentes de datos:")
    bala(doc, "CSV exportado de Google Forms — contiene respuestas reales de los estudiantes.")
    bala(doc, "Muestra aleatoria generada — 35 estudiantes sintéticos para demostración.")

    h(doc, "4.1  ¿Qué hace csv_importer.py?", 2)
    for item in [
        "Lee el archivo CSV con pandas (pd.read_csv).",
        "Mapea los nombres de columnas del formulario a los campos de Estudiante.",
        "Convierte texto a números: 'De 1 a 2 horas' → 1.5, 'Más de 20' → 22, etc.",
        "Valida rangos: calificaciones 0–10, asistencia 0–100, horas de trabajo 0–80.",
        "Registros con errores se omiten; los válidos se agregan al DataStore.",
        "Normaliza 'trabaja' desde columnas como 'Trabajas actualmente' → bool.",
    ]:
        bala(doc, item, "•")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 5. MÓDULOS ESTADÍSTICOS
    # ════════════════════════════════════════════════════════════════════
    h(doc, "5. Módulos Estadísticos", 1)

    h(doc, "5.1  Estadística Descriptiva (stats/descriptiva.py)", 2)
    p(doc, "Calcula todos los estadísticos de resumen para un conjunto de datos:")

    caja(doc, "Estadísticos calculados por ResumenDescriptivo", [
        ("n",               "Tamaño de la muestra"),
        ("media (x̄)",      "Σxᵢ / n"),
        ("mediana",         "Valor central al ordenar los datos"),
        ("moda",            "Valor más frecuente"),
        ("varianza (s²)",   "Σ(xᵢ − x̄)² / (n−1)  — varianza muestral, ddof=1"),
        ("desv. estándar (s)", "√s²"),
        ("coef. variación", "s / x̄ × 100%"),
        ("sesgo",           "scipy.stats.skew() — asimetría de la distribución"),
        ("curtosis",        "scipy.stats.kurtosis() — forma de las colas"),
        ("min, máx, rango", "Valores extremos y amplitud"),
        ("Q1, Q3, IQR",     "Cuartiles y rango intercuartílico"),
        ("IC 95%",          "x̄ ± t_(0.025, n−1) · (s/√n)  — intervalo de confianza"),
    ])

    p(doc, "Estos estadísticos se muestran en El Grimorio (tabla descriptiva) "
           "y en el Grimorio Avanzado con el procedimiento detallado.")

    h(doc, "5.2  Prueba de Hipótesis — Hipótesis B (stats/prueba_hipotesis.py)", 2)
    p(doc, "Este es el núcleo estadístico del proyecto. Contrasta si trabajar reduce el promedio.")

    p(doc, "Hipótesis planteadas:", bold=True)
    bala(doc, "H₀: μ_trabaja ≥ μ_no_trabaja  (trabajar NO afecta el promedio)")
    bala(doc, "H₁: μ_trabaja < μ_no_trabaja  (los que trabajan tienen MENOR promedio)")

    p(doc, "Tipo de prueba: t de Student para dos muestras independientes, "
           "una cola izquierda, varianzas iguales (pooled variance).", bold=True)

    caja(doc, "Pasos del procedimiento (también visibles en Grimorio Avanzado)", [
        ("1. Varianza combinada (sp²)",
         "sp² = [(n₁−1)s₁² + (n₂−1)s₂²] / (n₁+n₂−2)  ← promedio ponderado de varianzas"),
        ("2. Error estándar (SE)",
         "SE = sp · √(1/n₁ + 1/n₂)"),
        ("3. Estadístico t",
         "t = (x̄₁ − x̄₂) / SE"),
        ("4. Grados de libertad",
         "gl = n₁ + n₂ − 2"),
        ("5. p-valor",
         "P(T < t_obs | gl)  — calculado con scipy.stats.ttest_ind(alternative='less')"),
        ("6. Decisión",
         "Rechazar H₀ si p < α = 0.05  (o equivalente: t < t_crítico)"),
        ("7. IC 95% diferencia",
         "(x̄₁−x̄₂) ± t_(0.025, gl) · SE"),
        ("8. d de Cohen",
         "d = (x̄₁−x̄₂) / sp  — tamaño del efecto"),
    ])

    p(doc, "¿Por qué varianzas iguales (not Welch)?", bold=True)
    p(doc, "Asumimos que ambos grupos (trabaja / no trabaja) son parte de la misma población "
           "estudiantil y por tanto tienen variabilidad similar. Esta suposición da gl = n₁+n₂−2, "
           "más grados de libertad que Welch → prueba más potente.", indent=1, italic=True)

    h(doc, "5.3  Regresión Lineal Simple (stats/regresion_lineal.py)", 2)
    p(doc, "Modela la relación entre horas de estudio (x) y promedio final (y). "
           "Usa el método de mínimos cuadrados ordinarios (OLS).")

    caja(doc, "Fórmulas de regresión (visibles en Grimorio Avanzado Sección III)", [
        ("β₁ (pendiente)",  "β₁ = Σ(xᵢ−x̄)(yᵢ−ȳ) / Σ(xᵢ−x̄)²"),
        ("β₀ (intercepto)", "β₀ = ȳ − β₁·x̄"),
        ("Modelo",          "ŷ = β₀ + β₁·x"),
        ("R²",              "R² = 1 − SS_residual/SS_total  (% de varianza explicada)"),
    ])

    p(doc, "Interpretación de β₁: por cada hora adicional de estudio semanal, "
           "el promedio final cambia en β₁ puntos. Si β₁ > 0 → estudiar más mejora el promedio.", italic=True)

    h(doc, "5.4  Modelo de Predicción (stats/modelo_prediccion.py)", 2)
    p(doc, "La clase Oraculo predice la calificación esperada de un estudiante. "
           "Se usa en La Profecía y El Juicio Final.")

    caja(doc, "Fórmula de predicción del Oráculo", [
        ("base",      "μ del grupo (trabaja o no trabaja), calculada al 'entrenar'"),
        ("ajuste",    "(horas_estudio − 10) × 0.04  ← +0.04 pts por hora extra de estudio"),
        ("cal",       "clamp(base + ajuste, 0, 10)  ← acotada al rango válido"),
        ("P(aprobar)", "norm.sf(5.9, loc=cal, scale=σ_grupo)  ← distribución Normal acumulada"),
        ("nivel",     "cal ≥ 8.5 → Alto Rendimiento   cal ≥ 7.0 → Medio   resto → En Riesgo"),
    ])

    p(doc, "¿Por qué 0.04 por hora?", bold=True)
    p(doc, "Es un coeficiente empírico calibrado para que 10 horas de estudio no sumen ni resten "
           "(es el punto neutro), y que la diferencia entre 5 y 15 horas sea ±0.20 puntos, "
           "un efecto pequeño pero visible.", italic=True, indent=1)

    p(doc, "¿Por qué norm.sf(5.9, ...) para P(aprobar ≥ 6)?", bold=True)
    p(doc, "Se usa 5.9 en lugar de 6 para evitar que la probabilidad sea exactamente 50% "
           "cuando el promedio predicho es exactamente 6. La función sf (survival function) "
           "calcula 1 − Φ(z) = P(X > umbral).", italic=True, indent=1)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 6. PANTALLAS DE LA INTERFAZ
    # ════════════════════════════════════════════════════════════════════
    h(doc, "6. Pantallas de la Interfaz — Qué Hace Cada Una", 1)

    h(doc, "6.1  ⌂ Inicio (home_screen.py)", 2)
    p(doc, "Pantalla de bienvenida. Muestra: nombre del proyecto, hipótesis, instrucciones "
           "básicas de uso y los nombres del equipo. No tiene lógica estadística.")

    h(doc, "6.2  ⊕ Datos (data_entry.py)", 2)
    p(doc, "Permite cargar datos de dos formas:")
    bala(doc, "Importar CSV — abre un diálogo de archivo y llama a csv_importer.py.")
    bala(doc, "Generar muestra — llama a sample_generator.py para crear 35 estudiantes.")
    p(doc, "También muestra la tabla de estudiantes cargados y permite exportar a CSV nuevo. "
           "Cuando se cargan datos, el contador 'ALMAS REGISTRADAS' en la barra lateral se actualiza "
           "gracias al patrón Observer de DataStore.")

    h(doc, "6.3  ✦ La Profecía (prophecy_screen.py)", 2)
    p(doc, "Predicción individual de un estudiante. El usuario busca a un alumno por nombre "
           "o matrícula (buscador en vivo con Listbox), y el sistema muestra:")
    bala(doc, "Calificación predicha por el Oráculo.")
    bala(doc, "Nivel esperado (Alto / Medio / Riesgo).")
    bala(doc, "Probabilidad de aprobar como barra visual.")
    bala(doc, "Recomendaciones personalizadas según su situación laboral y horas de estudio.")
    p(doc, "El buscador filtra en tiempo real usando StringVar.trace_add('write', ...) "
           "que ejecuta la función de filtrado cada vez que el usuario escribe.", italic=True)

    h(doc, "6.4  ✧ El Grimorio (grimorio_screen.py)", 2)
    p(doc, "La pantalla más completa de visualización estadística. Contiene pestañas:")
    bala(doc, "Descriptiva: tabla con todos los estadísticos de resumen.")
    bala(doc, "Hipótesis: tabla de la prueba t de Student con resultado y conclusión.")
    bala(doc, "Histogramas: distribución de promedios y asistencia con curva normal superpuesta (área=1).")
    bala(doc, "Por carrera: boxplot comparativo y barras por nivel de estrés, género, estilo.")
    bala(doc, "Botón 'Exportar PDF': genera reporte de 4 páginas con matplotlib PdfPages.")
    p(doc, "Los histogramas usan density=True para que el área total bajo la curva sea 1, "
           "lo que permite superponer la curva de distribución normal correctamente.", italic=True)

    h(doc, "6.5  ∑ Grimorio Avanzado (grimorio_avanzado_screen.py)", 2)
    p(doc, "Pantalla pedagógica que muestra CADA fórmula sustituida con los datos reales del grupo cargado. "
           "Está dividida en 3 secciones:")
    bala(doc, "I. Estadística Descriptiva: media, varianza, desv. std., IC 95%.")
    bala(doc, "II. Prueba t de Student: varianza combinada, t, Cohen d, IC de la diferencia.")
    bala(doc, "III. Regresión Lineal Simple: β₀, β₁, R², interpretación.")
    p(doc, "Si el docente pregunta '¿cómo calcularon la varianza?', muestras esta pantalla "
           "y ves el número exacto con la sustitución completa.", italic=True, color="0D3B0D")

    h(doc, "6.6  ⚖ El Juicio Final (juicio_screen.py)", 2)
    p(doc, "Simula qué le pasaría al promedio de un estudiante si cambiara su situación laboral. "
           "Es la pantalla de análisis 'what-if' del sistema.")
    p(doc, "Cómo funciona la simulación:", bold=True)
    bala(doc, "Se selecciona un estudiante y se elige si simular que trabaja o no trabaja.")
    bala(doc, "El Oráculo calcula la calificación predicha para ese escenario (usando la media "
              "del grupo correspondiente + ajuste por horas de estudio).")
    bala(doc, "También calcula la predicción para su situación real.")
    bala(doc, "La pantalla muestra el PROCEDIMIENTO COMPLETO del cálculo: μ del grupo, "
              "ajuste por horas, fórmula, P(aprobar) con distribución Normal.")
    bala(doc, "La diferencia (Δ) muestra cuánto cambiaría su calificación al cambiar de grupo.")
    p(doc, "La gráfica de barras muestra la media de TODOS los que trabajan vs TODOS los que no, "
           "con la línea del promedio real del alumno seleccionado.", italic=True)

    h(doc, "6.7  ☽ Los Rituales (rituales_screen.py)", 2)
    p(doc, "Clasificación y recomendaciones automáticas para cada estudiante. Muestra:")
    bala(doc, "Tabla de clasificación: nivel de cada alumno (Alto / Medio / Riesgo).")
    bala(doc, "Alertas de baja asistencia (< 70%).")
    bala(doc, "Recomendaciones individuales: tutorías, grupos de estudio, plataformas.")
    bala(doc, "Opción de exportar la clasificación a CSV.")

    h(doc, "6.8  ? Guía de Uso (guia_screen.py)", 2)
    p(doc, "Manual integrado de la aplicación. Explica cada pantalla, la hipótesis estadística "
           "y las instrucciones de uso. No tiene lógica de cálculo.")

    h(doc, "6.9  ☆ Créditos (creditos_screen.py)", 2)
    p(doc, "Muestra la información del proyecto: institución, carrera, materia, docente, "
           "integrantes del equipo y referencias bibliográficas en formato APA 7ª edición.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 7. WIDGETS PERSONALIZADOS
    # ════════════════════════════════════════════════════════════════════
    h(doc, "7. Widgets Personalizados (ui/widgets.py)", 1)
    p(doc, "Componentes visuales reutilizables construidos sobre Tkinter:")

    caja(doc, "Widgets principales y su función", [
        ("GothicCard",        "Marco con borde y fondo oscuro — contenedor visual principal"),
        ("GothicButton",      "Botón estilizado con efecto hover y opción accent (dorado)"),
        ("ScrollableFrame",   "Frame con scrollbar vertical — permite contenido largo"),
        ("SectionTitle",      "Etiqueta decorativa para encabezados de sección"),
        ("TablaSimple",       "Tabla dinámica de 2 columnas con filas alternadas"),
        ("Medidor",           "Medidor semicircular tipo gauge — muestra calificación 0-10"),
        ("BarraProbabilidad", "Barra bicolor de progreso — P(aprobar) vs P(reprobar)"),
        ("color_nivel()",     "Función que devuelve el color según el nivel del alumno"),
    ])

    p(doc, "El Medidor (gauge) es un widget Canvas que dibuja arcos trigonométricos. "
           "Convierte la calificación 0-10 a un ángulo en la semicircunferencia.", italic=True)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 8. EXPORTACIÓN PDF
    # ════════════════════════════════════════════════════════════════════
    h(doc, "8. Exportación de Reporte PDF (stats/exportar_pdf.py)", 1)
    p(doc, "El botón 'Exportar PDF' en El Grimorio genera un documento de 4 páginas.")
    bala(doc, "Página 1 — Portada: título, fecha, número de estudiantes.")
    bala(doc, "Página 2 — Estadística Descriptiva: tabla con todos los estadísticos.")
    bala(doc, "Página 3 — Prueba de Hipótesis: tabla de la prueba t completa.")
    bala(doc, "Página 4 — Visualizaciones: histograma de promedios y barras de grupos.")

    p(doc, "Detalle técnico importante:", bold=True)
    p(doc, "Para generar el PDF sin abrir una ventana nueva, se cambia temporalmente el backend "
           "de matplotlib a 'Agg' (sin pantalla). Después de guardar el PDF, se restaura 'TkAgg'. "
           "Esto se hace con matplotlib.use('Agg') / matplotlib.use('TkAgg').", italic=True, indent=1)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 9. TECNOLOGÍAS Y BIBLIOTECAS
    # ════════════════════════════════════════════════════════════════════
    h(doc, "9. Tecnologías y Bibliotecas Utilizadas", 1)

    caja(doc, "Dependencias del proyecto (requirements.txt)", [
        ("Python 3.11+",      "Lenguaje de programación principal"),
        ("Tkinter",           "Interfaz gráfica — incluida en Python, sin instalación extra"),
        ("NumPy ≥ 1.24",      "Operaciones matemáticas vectorizadas — arrays, sumas, raíces"),
        ("SciPy ≥ 1.10",      "Estadística: ttest_ind, t.ppf, norm.sf, skew, kurtosis, shapiro"),
        ("pandas ≥ 2.0",      "Lectura de CSV (pd.read_csv) y construcción de DataFrames"),
        ("matplotlib ≥ 3.7",  "Gráficas, histogramas, boxplots, PDF (PdfPages), canvas Tkinter"),
        ("python-docx ≥ 1.1", "Generación de documentos Word (.docx) — este documento"),
    ])

    p(doc, "¿Por qué Tkinter y no PyQt o Electron?", bold=True)
    p(doc, "Tkinter está incluido en Python sin instalación adicional, lo que simplifica "
           "la distribución. Es suficiente para interfaces de escritorio educativas y "
           "permite integrar gráficas de matplotlib directamente.", italic=True, indent=1)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 10. PREGUNTAS FRECUENTES
    # ════════════════════════════════════════════════════════════════════
    h(doc, "10. Preguntas Frecuentes y Cómo Responderlas", 1)
    p(doc, "Esta sección cubre las preguntas más probables en una exposición de estadística aplicada.")

    preguntas = [
        ("¿Por qué usaron la prueba t de Student y no otra prueba?",
         "Porque comparamos las medias de dos grupos independientes (trabaja vs no trabaja) con muestras pequeñas "
         "y distribución aproximadamente normal. La prueba t de Student es la prueba paramétrica estándar para "
         "esta situación. Elegimos varianzas iguales (pooled) porque ambos grupos vienen de la misma población "
         "estudiantil, lo que da más potencia al test."),

        ("¿Qué significa el p-valor en su resultado?",
         "El p-valor es la probabilidad de obtener un estadístico t tan extremo (o más) por puro azar, "
         "SUPONIENDO que H₀ es verdadera. Si p < 0.05, hay menos del 5% de probabilidad de que la "
         "diferencia observada sea casualidad, por lo que rechazamos H₀ y concluimos que trabajar sí "
         "reduce el promedio estadísticamente."),

        ("¿Qué son los grados de libertad?",
         "gl = n₁ + n₂ − 2. Representan la cantidad de información 'libre' después de estimar los "
         "parámetros (las dos medias). Se restan 2 porque usamos la media de cada grupo. Más grados de "
         "libertad → la distribución t se aproxima más a la normal → prueba más precisa."),

        ("¿Qué es la d de Cohen?",
         "Es una medida del tamaño del efecto que no depende del tamaño de muestra. "
         "d = (x̄₁ − x̄₂) / sp. Escala: |d| < 0.2 pequeño, 0.2–0.5 mediano, 0.5–0.8 grande, > 0.8 muy grande. "
         "Complementa al p-valor: puedes tener p < 0.05 con d muy pequeño si n es grande."),

        ("¿El intervalo de confianza confirma la prueba?",
         "Sí. El IC 95% de la diferencia (μ₁−μ₂) es consistente con la prueba unilateral. "
         "Si el IC bilateral no cruza cero y es completamente negativo, confirma que μ_trabaja < μ_no_trabaja."),

        ("¿Cómo funciona exactamente la simulación del Juicio Final?",
         "El Oráculo estima la calificación así: (1) toma la media del grupo simulado como base; "
         "(2) ajusta con (horas_estudio − 10) × 0.04 para penalizar/premiar según cuánto estudia; "
         "(3) acota el resultado a [0, 10]; (4) calcula P(aprobar ≥ 6) con la distribución Normal. "
         "En la pantalla se ve cada paso del cálculo con los números reales del alumno."),

        ("¿Por qué los histogramas muestran density=True?",
         "Para poder superponer la curva de distribución normal. Con frequency=True las alturas "
         "dependen del número de datos y la escala no coincide con la función de densidad. "
         "Con density=True el área total es exactamente 1, y la curva Normal con misma media y "
         "desviación estándar tiene la misma escala que el histograma."),

        ("¿El CSV debe tener un formato específico?",
         "Sí. El importador espera el formato exportado de Google Forms con columnas específicas: "
         "nombre, matrícula, parciales, asistencia, trabaja, horas de trabajo y estudio. "
         "csv_importer.py intenta reconocer variaciones de nombre de columna (ej. 'Nombre completo' "
         "o 'Nombre') y convierte rangos de texto a número."),

        ("¿Qué pasa si no hay suficientes datos?",
         "Las pantallas verifican que haya al menos 2 estudiantes en cada grupo antes de calcular "
         "la prueba t. El Juicio Final requiere al menos 3 estudiantes. Si no se cumplen, "
         "se muestra un mensaje de advertencia en lugar de la pantalla."),

        ("¿Por qué usaron el Singleton para DataStore?",
         "Para garantizar que solo exista UNA lista de estudiantes en toda la aplicación. "
         "Si cada pantalla creara su propio objeto DataStore, los datos no se compartirían. "
         "Con el Singleton, DataStore.get() siempre devuelve la misma instancia con los mismos datos."),

        ("¿Cómo implementaron la regresión lineal?",
         "Con mínimos cuadrados ordinarios (OLS) implementado manualmente con NumPy. "
         "β₁ = Σ(xᵢ−x̄)(yᵢ−ȳ) / Σ(xᵢ−x̄)², β₀ = ȳ − β₁·x̄. También usamos scipy.stats.linregress "
         "para validar los resultados. El Grimorio Avanzado muestra cada suma con los datos reales."),

        ("¿Cuántos datos cargaron para la prueba?",
         "El CSV contiene los datos reales del formulario aplicado al grupo. También es posible "
         "generar 35 estudiantes de muestra para demostración. El sistema muestra el n de cada "
         "grupo en la tabla de la prueba t."),
    ]

    for i, (preg, resp) in enumerate(preguntas, 1):
        h(doc, f"P{i}: {preg}", 2)
        p(doc, resp, indent=0.5)
        doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # 11. GLOSARIO RÁPIDO
    # ════════════════════════════════════════════════════════════════════
    h(doc, "11. Glosario Rápido de Conceptos Estadísticos", 1)

    caja(doc, "Términos clave que debes dominar", [
        ("Hipótesis nula (H₀)",       "Lo que asumimos verdadero hasta tener evidencia en contra"),
        ("Hipótesis alternativa (H₁)", "Lo que queremos demostrar"),
        ("α (alfa)",                   "Nivel de significancia = probabilidad máxima de error tipo I (0.05 = 5%)"),
        ("p-valor",                    "Probabilidad de ver el resultado observado (o más extremo) si H₀ fuera cierta"),
        ("Error tipo I",               "Rechazar H₀ cuando es verdadera (falso positivo) — probabilidad = α"),
        ("Error tipo II",              "No rechazar H₀ cuando es falsa (falso negativo) — probabilidad = β"),
        ("Potencia (1−β)",             "Probabilidad de detectar el efecto si existe"),
        ("Estadístico t",              "Qué tan extrema es la diferencia observada en unidades de error estándar"),
        ("Grados de libertad (gl)",    "Información disponible para estimar variabilidad — n₁+n₂−2 para t de Student"),
        ("Varianza combinada (sp²)",   "Promedio ponderado de las varianzas de ambos grupos"),
        ("Intervalo de confianza",     "Rango dentro del cual el parámetro verdadero cae con 95% de probabilidad"),
        ("d de Cohen",                 "Tamaño del efecto estandarizado — no depende de n"),
        ("R²",                         "Coeficiente de determinación — % de varianza de y explicada por x"),
        ("Distribución Normal",        "Campana de Gauss — muchos fenómenos naturales la siguen"),
        ("Singleton",                  "Patrón de diseño que garantiza una sola instancia de una clase"),
        ("Observer",                   "Patrón donde suscriptores son notificados de cambios en un objeto"),
    ])

    # ════════════════════════════════════════════════════════════════════
    # PIE
    # ════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pie = pie.add_run(
        "\n\n¡Éxito en la exposición!\n\n"
        "Universidad Autónoma de Chiapas (UNACH)\n"
        "Probabilidad y Estadística · 4° \"O\"\n"
        f"Generado: {datetime.date.today().strftime('%d/%m/%Y')}"
    )
    r_pie.font.size   = Pt(12)
    r_pie.font.italic = True

    return doc


if __name__ == "__main__":
    doc  = crear_doc()
    ruta = (r"c:\Users\PC PRIDE WHITE WHALE\Downloads\EJ26"
            r"\Guia_Exposicion_Codigo.docx")
    doc.save(ruta)
    print(f"Documento creado: {ruta}")
